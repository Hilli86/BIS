"""
Thema Report - PDF/DOCX-Generierung für Themen
"""

from datetime import datetime
import os
import tempfile
import sys
from io import BytesIO
from flask import current_app
from docxtpl import DocxTemplate
from docx.shared import Mm, Pt
from utils.firmendaten import get_firmendaten
from utils.helpers import safe_get, format_schichtbuch_datum
from .pdf_export import convert_docx_to_pdf

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# Gleiche Bildtypen wie in modules/schichtbuch/services.get_thema_dateien_liste (+ gängige RAW-Viewer-Fälle)
_THEMA_BILD_ENDUNGEN = {
    '.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tif', '.tiff',
}


def _webp_als_temp_png(abs_path):
    """WebP für python-docx nach PNG konvertieren; Aufrufer löscht die Temp-Datei."""
    from PIL import Image

    im = Image.open(abs_path)
    if im.mode in ('RGBA', 'LA'):
        hintergrund = Image.new('RGB', im.size, (255, 255, 255))
        alpha = im.split()[-1] if im.mode == 'RGBA' else None
        hintergrund.paste(im, mask=alpha)
        im = hintergrund
    elif im.mode != 'RGB':
        im = im.convert('RGB')
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    tmp_name = tmp.name
    tmp.close()
    im.save(tmp_name, 'PNG')
    return tmp_name


def _thema_bilddateien_sortiert(thema_id, upload_ordner):
    """Liste (Anzeigename, absoluter Pfad) der Bilder im Thema-Upload, sortiert nach Dateiname."""
    if not upload_ordner:
        return []
    thema_ordner = os.path.join(upload_ordner, str(thema_id))
    if not os.path.isdir(thema_ordner):
        return []
    bilder = []
    for name in sorted(os.listdir(thema_ordner), key=str.lower):
        pfad = os.path.join(thema_ordner, name)
        if not os.path.isfile(pfad):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext in _THEMA_BILD_ENDUNGEN:
            bilder.append((name, os.path.abspath(pfad)))
    return bilder


def _fotos_an_dokument_anhaengen(doc_tpl, bilder, foto_breite_mm=140):
    """
    Nach docxtpl.render(): Bilder aus dem Thema-Ordner anhängen (keine Template-Platzhalter nötig).
    """
    if not bilder:
        return
    # Nicht get_docx() nutzen: docxtpl lädt bei is_rendered==True die Vorlage neu und verwirft render().
    dokument = doc_tpl.docx
    if dokument is None:
        return
    # Kein add_page_break(): erzeugt oft eine leere Seite (Vorlage/PDF-Konvertierung).
    titel = dokument.add_paragraph()
    titel.paragraph_format.space_before = Pt(18)
    titel_run = titel.add_run('Fotos zum Thema')
    titel_run.bold = True
    breite = Mm(foto_breite_mm)
    for dateiname, abs_pfad in bilder:
        einbinden_pfad = abs_pfad
        temp_png = None
        if os.path.splitext(abs_pfad)[1].lower() == '.webp':
            try:
                temp_png = _webp_als_temp_png(abs_pfad)
                einbinden_pfad = temp_png
            except Exception:
                einbinden_pfad = abs_pfad
                temp_png = None
        try:
            grafik_absatz = dokument.add_paragraph()
            grafik_absatz.add_run().add_picture(einbinden_pfad, width=breite)
            bildunterschrift = dokument.add_paragraph(dateiname)
            bildunterschrift.paragraph_format.space_after = Pt(10)
        except Exception:
            pass
        finally:
            if temp_png and os.path.isfile(temp_png):
                try:
                    os.unlink(temp_png)
                except OSError:
                    pass


def generate_thema_pdf(thema_id, conn):
    """
    Generiert ein PDF/DOCX für ein Thema.
    Gibt ein Tupel zurück: (content: bytes, filename: str, mimetype: str, is_pdf: bool)
    """
    # Thema-Informationen laden
    thema = conn.execute('''
        SELECT 
            t.ID,
            t.ErstelltAm,
            g.Bezeichnung AS Gewerk,
            b.Bezeichnung AS Bereich,
            s.Bezeichnung AS Status,
            s.Farbe AS StatusFarbe,
            a.Bezeichnung AS Abteilung
        FROM SchichtbuchThema t
        JOIN Gewerke g ON t.GewerkID = g.ID
        JOIN Bereich b ON g.BereichID = b.ID
        JOIN Status s ON t.StatusID = s.ID
        LEFT JOIN Abteilung a ON t.ErstellerAbteilungID = a.ID
        WHERE t.ID = ?
    ''', (thema_id,)).fetchone()
    
    if not thema:
        raise ValueError('Thema nicht gefunden.')
    
    # Sichtbare Abteilungen laden
    sichtbarkeiten = conn.execute('''
        SELECT a.Bezeichnung, a.ParentAbteilungID
        FROM SchichtbuchThemaSichtbarkeit sv
        JOIN Abteilung a ON sv.AbteilungID = a.ID
        WHERE sv.ThemaID = ?
        ORDER BY a.Sortierung, a.Bezeichnung
    ''', (thema_id,)).fetchall()

    # Optionale Zusatz-Gewerke laden (gleiche Sortierung wie in der UI)
    zusatz_gewerke = conn.execute('''
        SELECT g.Bezeichnung, b.Bezeichnung AS Bereich
        FROM SchichtbuchThemaGewerk tg
        JOIN Gewerke g ON tg.GewerkID = g.ID
        JOIN Bereich b ON g.BereichID = b.ID
        WHERE tg.ThemaID = ?
        ORDER BY b.Bezeichnung, g.Bezeichnung
    ''', (thema_id,)).fetchall()
    
    # Bemerkungen laden (chronologisch, älteste zuerst)
    bemerkungen = conn.execute('''
        SELECT 
            b.ID AS BemerkungID,
            b.Datum,
            m.Vorname,
            m.Nachname,
            b.Bemerkung,
            t.Bezeichnung AS Taetigkeit
        FROM SchichtbuchBemerkungen b
        JOIN Mitarbeiter m ON b.MitarbeiterID = m.ID
        LEFT JOIN Taetigkeit t ON b.TaetigkeitID = t.ID
        WHERE b.ThemaID = ? AND b.Gelöscht = 0
        ORDER BY b.Datum ASC
    ''', (thema_id,)).fetchall()
    
    # Ersatzteile für dieses Thema laden
    ersatzteile = conn.execute('''
        SELECT 
            l.ID AS BuchungsID,
            l.ErsatzteilID,
            l.Typ,
            l.Menge,
            l.Grund,
            l.Buchungsdatum,
            l.Bemerkung,
            l.Preis,
            l.Waehrung,
            e.Bestellnummer,
            e.Bezeichnung AS ErsatzteilBezeichnung,
            e.Einheit,
            m.Vorname || ' ' || m.Nachname AS VerwendetVon,
            k.Bezeichnung AS Kostenstelle
        FROM Lagerbuchung l
        JOIN Ersatzteil e ON l.ErsatzteilID = e.ID
        LEFT JOIN Mitarbeiter m ON l.VerwendetVonID = m.ID
        LEFT JOIN Kostenstelle k ON l.KostenstelleID = k.ID
        WHERE l.ThemaID = ?
        ORDER BY l.Buchungsdatum ASC
    ''', (thema_id,)).fetchall()
    
    # Firmendaten laden
    firmendaten = get_firmendaten()
    
    # Template-Pfad
    template_path = os.path.join(current_app.root_path, 'templates', 'reports', 'thema_template.docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError('Themenvorlage nicht gefunden.')
    
    # Template laden
    doc = DocxTemplate(template_path)
    
    # Datum formatieren
    thema_erstellt_am_formatiert = ''
    if thema['ErstelltAm']:
        try:
            thema_erstellt_am_formatiert = datetime.strptime(thema['ErstelltAm'], '%Y-%m-%d %H:%M:%S').strftime('%d.%m.%Y')
        except:
            thema_erstellt_am_formatiert = thema['ErstelltAm'][:10] if thema['ErstelltAm'] else ''
    
    export_datum = datetime.now()
    export_datum_formatiert = export_datum.strftime('%d.%m.%Y um %H:%M Uhr')
    
    # Sichtbarkeiten für Template vorbereiten
    sichtbarkeiten_liste = []
    for sichtbarkeit in sichtbarkeiten:
        sichtbarkeiten_liste.append({
            'bezeichnung': sichtbarkeit['Bezeichnung'] or ''
        })

    # Optionale Zusatz-Gewerke für Template vorbereiten
    zusatz_gewerke_liste = []
    for zg in zusatz_gewerke:
        zusatz_gewerke_liste.append({
            'bezeichnung': zg['Bezeichnung'] or '',
            'bereich': zg['Bereich'] or '',
        })
    thema_zusatz_gewerke_text = ', '.join(
        zg['bezeichnung'] for zg in zusatz_gewerke_liste if zg['bezeichnung']
    )
    
    # Bemerkungen für Template vorbereiten
    bemerkungen_liste = []
    for bemerkung in bemerkungen:
        datum_formatiert = format_schichtbuch_datum(bemerkung['Datum']) if bemerkung['Datum'] else ''
        
        mitarbeiter_name = f"{bemerkung['Vorname']} {bemerkung['Nachname']}"
        
        bemerkungen_liste.append({
            'datum': bemerkung['Datum'],
            'datum_formatiert': datum_formatiert,
            'mitarbeiter_vorname': bemerkung['Vorname'] or '',
            'mitarbeiter_nachname': bemerkung['Nachname'] or '',
            'mitarbeiter_name': mitarbeiter_name,
            'taetigkeit': bemerkung['Taetigkeit'] or '',
            'bemerkung': bemerkung['Bemerkung'] or ''
        })
    
    # Ersatzteile für Template vorbereiten
    ersatzteile_liste = []
    for ersatzteil in ersatzteile:
        datum_formatiert = ''
        if ersatzteil['Buchungsdatum']:
            try:
                datum_formatiert = datetime.strptime(ersatzteil['Buchungsdatum'], '%Y-%m-%d %H:%M:%S').strftime('%d.%m.%Y')
            except:
                datum_formatiert = ersatzteil['Buchungsdatum'][:10] if ersatzteil['Buchungsdatum'] else ''
        
        menge_val = ersatzteil['Menge'] if ersatzteil['Menge'] else 0
        einheit = ersatzteil['Einheit'] or 'Stück'
        if menge_val == int(menge_val):
            menge_mit_einheit = f"{int(menge_val)} {einheit}"
        else:
            menge_mit_einheit = f"{menge_val} {einheit}"
        
        preis_text = ''
        if ersatzteil['Preis']:
            preis_val = float(ersatzteil['Preis'])
            waehrung = ersatzteil['Waehrung'] or 'EUR'
            preis_text = f"{preis_val:.2f} {waehrung}"
        
        ersatzteile_liste.append({
            'datum': ersatzteil['Buchungsdatum'],
            'datum_formatiert': datum_formatiert,
            'ersatzteil_id': ersatzteil['ErsatzteilID'],
            'ersatzteil_bezeichnung': ersatzteil['ErsatzteilBezeichnung'] or '',
            'bestellnummer': ersatzteil['Bestellnummer'] or '',
            'typ': ersatzteil['Typ'] or '',
            'menge': menge_val,
            'einheit': einheit,
            'menge_mit_einheit': menge_mit_einheit,
            'verwendet_von': ersatzteil['VerwendetVon'] or '',
            'kostenstelle': ersatzteil['Kostenstelle'] or '',
            'preis': preis_text,
            'waehrung': ersatzteil['Waehrung'] or 'EUR'
        })
    
    # Footer-Informationen zusammenstellen
    footer_lines = []
    if firmendaten:
        if safe_get(firmendaten, 'Geschaeftsfuehrer'):
            footer_lines.append(f"Geschäftsführer: {firmendaten['Geschaeftsfuehrer']}")
        if safe_get(firmendaten, 'UStIdNr'):
            footer_lines.append(f"UStIdNr.: {firmendaten['UStIdNr']}")
        if safe_get(firmendaten, 'Steuernummer'):
            footer_lines.append(f"Steuernr.: {firmendaten['Steuernummer']}")
        if safe_get(firmendaten, 'Telefon'):
            footer_lines.append(f"Telefon: {firmendaten['Telefon']}")
        bank_name = safe_get(firmendaten, 'BankName')
        iban = safe_get(firmendaten, 'IBAN')
        if bank_name and iban:
            footer_lines.append(f"Bankverbindung: {firmendaten['BankName']}")
            footer_lines.append(f"IBAN: {firmendaten['IBAN']}")
            bic = safe_get(firmendaten, 'BIC')
            if bic:
                footer_lines.append(f"BIC: {firmendaten['BIC']}")
    footer_text = ' | '.join(footer_lines) if footer_lines else ''
    
    # Kontext für Template
    context = {
        # Thema-Daten
        'thema_id': thema['ID'],
        'thema_bereich': thema['Bereich'] or '',
        'thema_gewerk': thema['Gewerk'] or '',
        'thema_zusatz_gewerke': thema_zusatz_gewerke_text,
        'thema_status': thema['Status'] or '',
        'thema_status_farbe': thema['StatusFarbe'] or '',
        'thema_abteilung': thema['Abteilung'] or '',
        'thema_erstellt_am': thema['ErstelltAm'],
        'thema_erstellt_am_formatiert': thema_erstellt_am_formatiert,
        
        # Sichtbarkeiten
        'sichtbarkeiten': sichtbarkeiten_liste,
        'hat_sichtbarkeiten': len(sichtbarkeiten_liste) > 0,

        # Optionale Zusatz-Gewerke
        'zusatz_gewerke': zusatz_gewerke_liste,
        'hat_zusatz_gewerke': len(zusatz_gewerke_liste) > 0,
        
        # Bemerkungen
        'bemerkungen': bemerkungen_liste,
        'hat_bemerkungen': len(bemerkungen_liste) > 0,
        
        # Ersatzteile
        'ersatzteile': ersatzteile_liste,
        'hat_ersatzteile': len(ersatzteile_liste) > 0,
        
        # Firmendaten
        'firmenname': safe_get(firmendaten, 'Firmenname', '') if firmendaten else '',
        'firmenstrasse': safe_get(firmendaten, 'Strasse', '') if firmendaten else '',
        'firmenplz': safe_get(firmendaten, 'PLZ', '') if firmendaten else '',
        'firmenort': safe_get(firmendaten, 'Ort', '') if firmendaten else '',
        'firmenplz_ort': f"{safe_get(firmendaten, 'PLZ', '')} {safe_get(firmendaten, 'Ort', '')}".strip() if firmendaten else '',
        'firmen_telefon': safe_get(firmendaten, 'Telefon', '') if firmendaten else '',
        'firmen_website': safe_get(firmendaten, 'Website', '') if firmendaten else '',
        'firmen_email': safe_get(firmendaten, 'Email', '') if firmendaten else '',
        'footer': footer_text,
        
        # Metadaten
        'export_datum': export_datum.isoformat(),
        'export_datum_formatiert': export_datum_formatiert,
    }
    
    # Template rendern
    doc.render(context)

    upload_ordner = current_app.config.get('SCHICHTBUCH_UPLOAD_FOLDER')
    thema_bilder = _thema_bilddateien_sortiert(thema_id, upload_ordner)
    _fotos_an_dokument_anhaengen(doc, thema_bilder)
    
    # Als PDF konvertieren oder DOCX zurückgeben
    if DOCX2PDF_AVAILABLE:
        # PDF-Konvertierung versuchen
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Temporäre DOCX-Datei erstellen
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            tmp_docx.write(buffer.getvalue())
            tmp_docx_path = tmp_docx.name
        
        # PDF erstellen
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
            tmp_pdf_path = tmp_pdf.name
        
        try:
            # PDF-Konvertierung
            if convert_docx_to_pdf(tmp_docx_path, tmp_pdf_path):
                # PDF lesen
                with open(tmp_pdf_path, 'rb') as f:
                    pdf_content = f.read()
                
                # Temporäre Dateien löschen
                os.unlink(tmp_docx_path)
                os.unlink(tmp_pdf_path)
                
                filename = f"Thema_{thema_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
                return (pdf_content, filename, 'application/pdf', True)
        except Exception as e:
            # Temporäre Dateien aufräumen
            if os.path.exists(tmp_docx_path):
                try:
                    os.unlink(tmp_docx_path)
                except:
                    pass
            if os.path.exists(tmp_pdf_path):
                try:
                    os.unlink(tmp_pdf_path)
                except:
                    pass
            
            # COM aufräumen (Windows)
            if sys.platform == 'win32':
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except:
                    pass
    
    # Fallback: DOCX zurückgeben
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"Thema_{thema_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    return (buffer.getvalue(), filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', False)

